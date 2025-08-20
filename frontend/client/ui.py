import streamlit as st
import requests
import tempfile
import traceback
import xml.etree.ElementTree as ET
import json
import time
import os
import shutil
import zipfile
from docx import Document as DocxDocument
from fpdf import FPDF
from io import BytesIO
from typing import List, Dict, Any, Optional

# Configuration
BACKEND_URL = os.getenv("BACKEND_URL", "http://127.0.0.1:5001")
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB per file
SUPPORTED_EXTENSIONS = ['.pdf', '.docx', '.txt', '.csv', '.xlsx', '.png', '.jpg', '.jpeg']
CONFIG_EXTENSIONS = ['yaml', 'yml', 'json']


st.title("Document Analyzer")
st.markdown("""
    **Advanced document analysis workflow:**  
    1. Upload configuration files  
    2. Upload documents (multiple files or zipped folders)  
""")




def init_session_state():
    """Initialize session state variables"""
    defaults = {
        'session_id': None,
        'config_uploaded': False,
        'uploaded_configs': [],
        'temp_dir': None,
        'extracted_files': [],
        'analysis_complete': False,
        'show_results': False,
        'extraction_results': {},
        'text_sample': ''
    }

    for key, default_value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default_value

def cleanup_temp_directory():
    """Clean up temporary directory"""
    if st.session_state.temp_dir and os.path.exists(st.session_state.temp_dir):
        try:
            shutil.rmtree(st.session_state.temp_dir)
            st.session_state.temp_dir = None
        except Exception as e:
            st.warning(f"Could not clean up temporary directory: {str(e)}")

def reset_session():
    """Reset all session state variables"""
    cleanup_temp_directory()
    for key in list(st.session_state.keys()):
        if key.startswith('form_'):  # Keep form keys
            continue
        del st.session_state[key]
    init_session_state()

def validate_file_size(file) -> bool:
    """Check if file size is within limits"""
    if hasattr(file, 'size') and file.size > MAX_FILE_SIZE:
        return False
    return True

def extract_zip_files(zip_file) -> List[str]:
    """Extract ZIP file and return list of valid document paths"""
    try:
        if not st.session_state.temp_dir:
            st.session_state.temp_dir = tempfile.mkdtemp()
        
        # Save uploaded ZIP file
        zip_path = os.path.join(st.session_state.temp_dir, "uploaded.zip")
        with open(zip_path, "wb") as f:
            f.write(zip_file.getvalue())
        
        # Extract ZIP file
        extract_dir = os.path.join(st.session_state.temp_dir, "extracted")
        os.makedirs(extract_dir, exist_ok=True)
        
        extracted_files = []
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            # Check for zip bombs or excessive files
            file_list = zip_ref.namelist()
            if len(file_list) > 100:  # Limit number of files
                st.warning("ZIP contains too many files (max 100)")
                return []
            
            for file_info in zip_ref.infolist():
                # Skip directories and hidden files
                if file_info.is_dir() or file_info.filename.startswith('.'):
                    continue
                
                # Check file extension
                file_ext = os.path.splitext(file_info.filename.lower())[1]
                if file_ext in SUPPORTED_EXTENSIONS:
                    # Check file size
                    if file_info.file_size > MAX_FILE_SIZE:
                        st.warning(f"Skipping {file_info.filename} - file too large")
                        continue
                    
                    # Extract file
                    zip_ref.extract(file_info, extract_dir)
                    extracted_files.append(os.path.join(extract_dir, file_info.filename))
        
        return extracted_files
        
    except zipfile.BadZipFile:
        st.error("Invalid ZIP file")
        return []
    except Exception as e:
        st.error(f"Error extracting ZIP: {str(e)}")
        return []

def upload_config_files(config_files) -> bool:
    """Upload configuration files to backend"""
    try:
        session_ids = []
        uploaded_configs = []
        
        for config_file in config_files:
            if not validate_file_size(config_file):
                st.warning(f"Config file {config_file.name} is too large")
                continue
                
            files = {"config_file": (config_file.name, config_file.getvalue())}
            
            with st.status(f"Uploading {config_file.name}..."):
                response = requests.post(
                    f"{BACKEND_URL}/upload_config",
                    files=files,
                    timeout=30
                )
                
                if response.status_code == 200:
                    session_id = response.json().get("session_id")
                    if session_id:
                        session_ids.append(session_id)
                        uploaded_configs.append({
                            "name": config_file.name,
                            "session_id": session_id,
                            "content": config_file.getvalue()
                        })
                        st.success(f" {config_file.name}")
                    else:
                        st.error(f" {config_file.name} - No session ID returned")
                else:
                    st.error(f" {config_file.name} - {response.text}")
        
        if session_ids:
            st.session_state.session_id = session_ids[0]
            st.session_state.uploaded_configs = uploaded_configs
            st.session_state.config_uploaded = True
            return True
        
        return False
        
    except requests.exceptions.ConnectionError:
        st.error(" Backend connection failed. Is the server running?")
        return False
    except requests.exceptions.Timeout:
        st.error(" Request timeout. Server may be overloaded.")
        return False
    except Exception as e:
        st.error(f" Error processing config files: {str(e)}")
        return False

def process_documents(files_data: List[tuple], session_id: str) -> bool:
    """Process documents with backend"""
    try:
        with st.status("Analyzing documents...") as status:
            response = requests.post(
                f"{BACKEND_URL}/upload_documents",
                files=files_data,
                data={"session_id": session_id},
                timeout=300  # 5 minutes for large documents
            )
            
            if response.status_code == 200:
                result = response.json()
                st.session_state.extraction_results = result.get("data", {})
                st.session_state.text_sample = result.get("text_sample", "")
                st.session_state.analysis_complete = True
                st.session_state.show_results = False
                status.update(label=" Analysis complete!", state="complete")
                return True
            else:
                st.error(f" Backend error: {response.text}")
                return False
                
    except requests.exceptions.ConnectionError:
        st.error(" Backend connection failed during analysis")
        return False
    except requests.exceptions.Timeout:
        st.error(" Analysis timeout. Try with fewer or smaller files.")
        return False
    except Exception as e:
        st.error(f" Analysis failed: {str(e)}")
        return False


# Initialize session state
init_session_state()

# Sidebar
with st.sidebar:
    st.header("Workflow")
    st.markdown("""
    1. Upload config files (YAML/JSON)
    2. Upload documents (files/zip)
    3. Analyze documents
    4. Download results
    """)
    
    if st.session_state.config_uploaded:
        st.success(f" {len(st.session_state.uploaded_configs)} config file(s) uploaded")
        if st.button(" Clear Session"):
            reset_session()
            st.rerun()

# Step 1: Multiple Config File Upload
st.write("")
st.subheader("Step 1: Upload Configuration Files")
st.markdown("### Sample YAML file ")
st.code("""
  fields:
    - name: "John Doe"
      keywords: ["insurance", "health"]
      response_type: "detailed"
    - name: "Jane Smith"
      keywords: ["auto", "accident"]
      response_type: "concise"
""", language="yaml")
st.write("")
st.write("")

if not st.session_state.config_uploaded:
    config_files = st.file_uploader(
        "Choose configuration files (YAML or JSON)",
        type=CONFIG_EXTENSIONS,
        accept_multiple_files=True,
        key="config_upload"
    )
    
    if config_files:
        st.info(f"Selected {len(config_files)} config file(s)")
        
        if st.button(" Upload Configurations", type="primary"):
            if upload_config_files(config_files):
                st.rerun()
else:
    st.success(" Configuration files uploaded successfully!")
    
    # Show uploaded configs
    with st.expander("View uploaded configurations"):
        for i, config in enumerate(st.session_state.uploaded_configs):
            st.write(f"**{i+1}.** {config['name']} (Session: {config['session_id'][:8]}...)")

# Step 2: Document Upload and Processing
if st.session_state.config_uploaded:
    st.subheader("Step 2: Upload Documents")
    
    # optional keyword input
    st.caption("Add keywords through text (Optional)")
    additional_keywords = st.text_input(
        "Extra keywords (comma separated):",
        placeholder="e.g., contract, agreement, termination",
        help="Add additional terms to focus the analysis on specific concepts"
    )

    # Upload method selection
    upload_option = st.radio(
        "Upload method:",
        ["Individual Files", "ZIP Archive"],
        horizontal=True
    )
    
    # Initialize variables
    uploaded_files = None
    files_ready = False
    
    if upload_option == "Individual Files":
        uploaded_files = st.file_uploader(
            "Choose documents to analyze",
            type=["pdf", "docx", "txt", "xlsx", "csv", "png", "jpg", "jpeg"],
            accept_multiple_files=True,
            key="doc_upload"
        )
        
        if uploaded_files:
            # Validate file sizes
            valid_files = []
            for file in uploaded_files:
                if validate_file_size(file):
                    valid_files.append(file)
                else:
                    st.warning(f"Skipping {file.name} - file too large (max {MAX_FILE_SIZE//1024//1024}MB)")
            
            uploaded_files = valid_files
            if uploaded_files:
                st.success(f" Selected {len(uploaded_files)} valid file(s)")
                files_ready = True
    
    elif upload_option == "ZIP Archive":
        zip_file = st.file_uploader(
            "Upload ZIP archive containing documents",
            type=["zip"],
            key="zip_upload"
        )
        
        if zip_file:
            if not validate_file_size(zip_file):
                st.error(f"ZIP file too large (max {MAX_FILE_SIZE//1024//1024}MB)")
            else:
                with st.spinner(" Extracting ZIP file..."):
                    extracted_files = extract_zip_files(zip_file)
                    
                if extracted_files:
                    st.session_state.extracted_files = extracted_files
                    st.success(f" Extracted {len(extracted_files)} document(s)")
                    files_ready = True
                else:
                    st.error("No valid documents found in ZIP archive")

    # Process documents when ready

    if files_ready:
        # Config selection (if multiple configs)
        selected_session_id = st.session_state.session_id
        if len(st.session_state.uploaded_configs) > 1:
            selected_config = st.selectbox(
                "Select configuration to use:",
                [cfg["name"] for cfg in st.session_state.uploaded_configs],
                index=0
            )
            selected_session_id = next(
                cfg["session_id"] for cfg in st.session_state.uploaded_configs
                if cfg["name"] == selected_config
            )
        
        if st.button(" Analyze Documents", type="primary"):
            # Prepare files for upload
            files_data = []
            
            try:
                if upload_option == "Individual Files" and uploaded_files:
                    for file in uploaded_files:
                        files_data.append(("document_files", (file.name, file.getvalue())))
                        
                elif upload_option == "ZIP Archive" and st.session_state.extracted_files:
                    for file_path in st.session_state.extracted_files:
                        if os.path.exists(file_path):
                            with open(file_path, 'rb') as f:
                                files_data.append(("document_files", (os.path.basename(file_path), f.read())))
                
                if files_data:
                    # Modified to include keywords in the request
                    data = {"session_id": selected_session_id}
                    if additional_keywords:
                        data["keywords"] = additional_keywords
                    
                    response = requests.post(
                        f"{BACKEND_URL}/upload_documents",
                        files=files_data,
                        data=data,  # Now includes optional keywords
                        timeout=300
                    )
                    
                    if response.status_code == 200:
                        result = response.json()
                        st.session_state.extraction_results = result.get("data", {})
                        st.session_state.text_sample = result.get("text_sample", "")
                        st.session_state.analysis_complete = True
                        st.session_state.show_results = False
                        st.rerun()
                    else:
                        st.error(f" Backend error: {response.text}")
                else:
                    st.error("No files to process")
                    
            except Exception as e:
                st.error(f"Error processing documents: {str(e)}")

# Step 3: Results Section
if st.session_state.analysis_complete:
    st.subheader("Step 3: View Results")
    
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button(" Show Results") and not st.session_state.show_results:
            st.session_state.show_results = True
            st.rerun()
    
    with col2:
        if st.session_state.show_results and st.button(" Hide Results"):
            st.session_state.show_results = False
            st.rerun()
    
    if st.session_state.show_results:
        # Display results
        st.subheader(" Analysis Results")
        
        # Text sample preview
        if st.session_state.text_sample:
            with st.expander(" View Document Sample"):
                st.text_area(
                    "Sample Text", 
                    st.session_state.text_sample, 
                    height=200,
                    disabled=True
                )
        
        # Output format selection
        output_format = st.radio(
            "Select output format:",
            ["JSON", "Text", "XML"],
            horizontal=True,
            key="output_format"
        )
        
        # Display results in selected format
        try:
            if output_format == "JSON":
                st.json(st.session_state.extraction_results)
                
            elif output_format == "Text":
                if "results" in st.session_state.extraction_results:
                    for i, item in enumerate(st.session_state.extraction_results["results"], 1):
                        st.markdown(f"### Result {i}: {item.get('field', 'N/A')}")
                        st.markdown(f"**Type:** {item.get('type', 'N/A')}")
                        st.markdown(f"**Confidence:** {item.get('confidence', 'N/A')}")
                        st.markdown("**Value:**")
                        st.write(item.get('value', 'No value extracted'))
                        st.divider()
                else:
                    st.info("No results to display")
                    
            elif output_format == "XML":
                root = ET.Element("AnalysisResults")
                if "results" in st.session_state.extraction_results:
                    for item in st.session_state.extraction_results["results"]:
                        result_elem = ET.SubElement(root, "Result")
                        ET.SubElement(result_elem, "Field").text = str(item.get('field', ''))
                        ET.SubElement(result_elem, "Type").text = str(item.get('type', ''))
                        ET.SubElement(result_elem, "Confidence").text = str(item.get('confidence', ''))
                        ET.SubElement(result_elem, "Value").text = str(item.get('value', ''))
                
                xml_str = ET.tostring(root, encoding='unicode')
                st.code(xml_str, language="xml")
            
                
        except Exception as e:
            st.error(f"Error displaying results: {str(e)}")
        
        # Download section
        st.subheader(" Download Results")
        d1_col1, d1_col2 = st.columns(2)
        file_data = None
        filename = ""
        mime = ""

        try:

            with d1_col1:
                if st.button("Download as JSON"):
                    file_data = json.dumps(st.session_state.extraction_results, indent=2).encode('utf-8')
                    filename = "extraction_results.json"
                    mime = "application/json"

                if st.button("Download as Text"):
                    text_content = ""
                    if "results" in st.session_state.extraction_results:
                        for i, item in enumerate(st.session_state.extraction_results["results"], 1):
                            text_content += f"Result {i}:\n"
                            text_content += f"Field: {item.get('field', 'N/A')}\n"
                            text_content += f"Type: {item.get('type', 'N/A')}\n"
                            text_content += f"Confidence: {item.get('confidence', 'N/A')}\n"
                            text_content += f"Value: {item.get('value', 'N/A')}\n"
                            text_content += "-" * 50 + "\n\n"
                    file_data = text_content.encode('utf-8')
                    filename = "extraction_results.txt"
                    mime = "text/plain"

                if st.button("Download as XML"):
                    root = ET.Element("AnalysisResults")
                    if "results" in st.session_state.extraction_results:
                        for item in st.session_state.extraction_results["results"]:
                            result_elem = ET.SubElement(root, "Result")
                            ET.SubElement(result_elem, "Field").text = str(item.get('field', ''))
                            ET.SubElement(result_elem, "Type").text = str(item.get('type', ''))
                            ET.SubElement(result_elem, "Confidence").text = str(item.get('confidence', ''))
                            ET.SubElement(result_elem, "Value").text = str(item.get('value', ''))
                    
                    xml_str = ET.tostring(root, encoding='unicode')
                    file_data = xml_str.encode('utf-8')
                    filename = "extraction_results.xml"
                    mime = "application/xml"
                    
                if st.button("Download as DOCX"):
                    result_data = st.session_state.get("extraction_results", {})
                    output = BytesIO()
                    doc = DocxDocument()
                    doc.add_heading("Extracted Data", level=1)
                    for item in result_data["results"]:
                        doc.add_paragraph(f"{item['field']}: {item['value']}")  # Key-value pair
                    doc.save(output)
                    output.seek(0)
                    file_data = output.getvalue()
                    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    filename = "extraction_results.docx"

                if st.button("Download as PDF"):
                    result_data = st.session_state.get("extraction_results", {})
                    output = BytesIO()
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)

                    for item in result_data["results"]:
                        field = item['field']
                        value = item['value']
                        # Safely encode for PDF
                        safe_value = str(value).encode('latin-1', 'replace').decode('latin-1')
                        pdf.multi_cell(0, 10, f"{field}: {safe_value}")

                    # Write PDF content to BytesIO
                    pdf_output = pdf.output(dest='S').encode('latin-1')
                    output.write(pdf_output)
                    output.seek(0)
                    file_data = output.getvalue()
                    mime = "application/pdf"
                    filename = "extraction_results.pdf"

            with d1_col2:
                if file_data:
                    st.download_button("Click to save", file_data, filename, mime=mime)
            
                
        except Exception as e:
            st.error(f"Error preparing downloads: {str(e)}")

# Cleanup on exit
if st.session_state.get('temp_dir') and os.path.exists(st.session_state.temp_dir):
    try:
        # Only cleanup if not actively using temp files
        if not st.session_state.extracted_files:
            shutil.rmtree(st.session_state.temp_dir)
            st.session_state.temp_dir = None
    except:
        pass

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center'>
    <p><em>Document Analyzer - Powered by AI Processing</em></p>
    <p><small>Backend: {}</small></p>
</div>
""".format(BACKEND_URL), unsafe_allow_html=True)
